#!/usr/bin/env python3
"""
Comprehensive Figure Integration Script v2

Improvements:
1. Removes duplicate figure caption text (keeps only figure element)
2. Adds ALL figure references (placeholder for missing images)
3. Smaller figures with click-to-expand functionality
4. Better matching logic
"""

import json
import os
import re
from pathlib import Path
from docx import Document

# Paths
SOURCE_DIR = 'project_assets/images/dissertation/By_Chapter'
PUBLIC_DIR = 'public/assets/dissertation/figures'
EXTRACTED_DIR = 'project_assets/docs/dissertation/extracted'
DOCX_PATH = 'project_assets/docs/dissertation/Schurtz_Dissertation_Final.docx'

def extract_all_figure_captions():
    """Extract complete figure list with captions from Word doc"""
    doc = Document(DOCX_PATH)
    
    figures = {}
    
    # Figure list is in paragraphs 201-250
    for i in range(201, 260):
        if i >= len(doc.paragraphs):
            break
        
        text = doc.paragraphs[i].text.strip()
        
        # Match "Figure X.Y: Caption"
        match = re.match(r'Figure\s+(\d+\.\d+):\s+(.+?)(?:\s+\d+\s*$|$)', text)
        if match:
            fig_num = match.group(1)
            caption = match.group(2).strip()
            
            # Remove trailing page numbers
            caption = re.sub(r'\s+\d+\s*$', '', caption)
            
            chapter = fig_num.split('.')[0]
            if chapter not in figures:
                figures[chapter] = {}
            
            figures[chapter][fig_num] = caption
    
    return figures

def find_best_image_match(chapter, fig_ref, caption, available_images):
    """Try to find best matching image for a figure"""
    
    if chapter not in available_images:
        return None
    
    chapter_images = available_images[chapter]
    
    # Extract key terms from caption
    caption_lower = caption.lower()
    
    # Try different matching strategies
    for img in chapter_images:
        filename_lower = img['filename'].lower()
        
        # Strategy 1: Figure number in filename
        fig_patterns = [
            fig_ref.replace('.', '-'),
            fig_ref.replace('.', '_'),
            fig_ref.replace('.', ''),
        ]
        
        for pattern in fig_patterns:
            if pattern in filename_lower.replace(' ', '').replace('_', '').replace('-', ''):
                return img
        
        # Strategy 2: Key caption words in filename
        key_words = []
        if 'map' in caption_lower:
            key_words.append('map')
        if 'plan' in caption_lower or 'building' in caption_lower:
            key_words.extend(['plan', 'building'])
        if 'pithos' in caption_lower or 'vessel' in caption_lower:
            key_words.extend(['pithos', 'vessel', 'drum'])
        if 'spearhead' in caption_lower:
            key_words.append('spearhead')
        if 'relief' in caption_lower:
            key_words.append('relief')
        if 'khorsabad' in caption_lower or 'musasir' in caption_lower or 'muṣaṣir' in caption_lower:
            key_words.extend(['khorsabad', 'musasir'])
        if 'mudjesir' in caption_lower:
            key_words.append('mudjesir')
        if 'corona' in caption_lower:
            key_words.append('corona')
        if 'maxar' in caption_lower or 'satellite' in caption_lower:
            key_words.extend(['maxar', 'satellite'])
        if 'section' in caption_lower:
            key_words.append('section')
        if 'burial' in caption_lower:
            key_words.append('burial')
        if 'bead' in caption_lower:
            key_words.append('bead')
        if 'topzawa' in caption_lower or 'gund-i' in caption_lower:
            key_words.extend(['topzawa', 'gund'])
        
        # Check if multiple key words match
        matches = sum(1 for word in key_words if word in filename_lower)
        if matches >= 2:
            return img
    
    return None

def scan_available_images():
    """Scan all available images"""
    images = {}
    
    for chapter_dir in Path(SOURCE_DIR).glob('Chapter *'):
        chapter_num = chapter_dir.name.split()[-1]
        images[chapter_num] = []
        
        # Get all image files
        for ext in ['*.jpg', '*.jpeg', '*.png', '*.pdf']:
            for img_path in chapter_dir.rglob(ext):
                if not img_path.name.startswith('.'):
                    rel_path = img_path.relative_to(chapter_dir)
                    images[chapter_num].append({
                        'filename': img_path.name,
                        'path': str(img_path),
                        'ext': img_path.suffix.lower(),
                        'subfolder': str(rel_path.parent) if rel_path.parent != Path('.') else None
                    })
    
    return images

def create_figure_html(chapter, fig_ref, caption, img_path=None, is_placeholder=False):
    """Create HTML for a figure (with or without image)"""
    
    if is_placeholder:
        return f'''
<figure class="my-8 p-6 border border-dashed border-gray-600 bg-gray-900/20">
  <div class="text-center py-12 text-gray-500">
    <svg class="mx-auto h-12 w-12 mb-4" fill="none" stroke="currentColor" viewBox="0 0 24 24">
      <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M4 16l4.586-4.586a2 2 0 012.828 0L16 16m-2-2l1.586-1.586a2 2 0 012.828 0L20 14m-6-6h.01M6 20h12a2 2 0 002-2V6a2 2 0 00-2-2H6a2 2 0 00-2 2v12a2 2 0 002 2z" />
    </svg>
    <p class="text-sm font-mono">Figure {fig_ref} not yet available</p>
  </div>
  <figcaption class="text-sm text-gray-500 mt-2">Figure {fig_ref}: {caption}</figcaption>
</figure>
'''
    
    # Real image with click-to-expand
    return f'''
<figure class="my-8">
  <a href="{img_path}" target="_blank" class="block cursor-zoom-in hover:opacity-90 transition-opacity">
    <img src="{img_path}" alt="Figure {fig_ref}: {caption}" class="w-full max-w-2xl mx-auto border border-white/10" loading="lazy" />
  </a>
  <figcaption class="text-sm text-gray-500 mt-2 text-center">Figure {fig_ref}: {caption} <span class="text-xs text-gray-600">(click to enlarge)</span></figcaption>
</figure>
'''

def process_chapter(chapter_num, figures, available_images, public_dir):
    """Process a single chapter - add all figures, remove duplicates"""
    
    md_file = Path(EXTRACTED_DIR) / f"chapter-{chapter_num}.md"
    
    if not md_file.exists():
        return None
    
    with open(md_file, 'r') as f:
        content = f.read()
    
    if chapter_num not in figures:
        return None
    
    chapter_figures = figures[chapter_num]
    stats = {'added': 0, 'placeholder': 0, 'removed_dups': 0}
    
    # Process each figure
    for fig_ref, caption in sorted(chapter_figures.items()):
        # Find or create image
        img_match = find_best_image_match(chapter_num, fig_ref, caption, available_images)
        
        if img_match:
            # Copy to public folder
            dest_filename = f"figure-{chapter_num}-{fig_ref.replace('.', '-')}{img_match['ext']}"
            dest_path = Path(public_dir) / dest_filename
            
            if not dest_path.exists():
                import shutil
                shutil.copy2(img_match['path'], dest_path)
            
            img_url = f"/assets/dissertation/figures/{dest_filename}"
            figure_html = create_figure_html(chapter_num, fig_ref, caption, img_url, False)
            stats['added'] += 1
        else:
            # Placeholder
            figure_html = create_figure_html(chapter_num, fig_ref, caption, None, True)
            stats['placeholder'] += 1
        
        # Find the figure reference in text
        # Pattern matches standalone "Figure X.Y" text (not already in HTML)
        pattern = rf'(?<!\<figcaption)Figure {re.escape(fig_ref)}(?::|\s)'
        
        matches = list(re.finditer(pattern, content, re.IGNORECASE))
        
        if matches:
            # Take first match
            match = matches[0]
            start = match.start()
            
            # Find end of paragraph/line containing this reference
            line_end = content.find('\n', start)
            if line_end == -1:
                line_end = len(content)
            
            # Check if there's already a figure tag nearby
            check_before = content[max(0, start-200):start]
            check_after = content[start:min(len(content), start+200)]
            
            if '<figure' not in check_before and '<figure' not in check_after:
                # Remove the caption line if it exists
                caption_pattern = rf'^Figure {re.escape(fig_ref)}:.*$'
                caption_line_match = re.search(caption_pattern, content[start:line_end], re.MULTILINE)
                
                if caption_line_match:
                    # Remove this line entirely
                    line_start = start + caption_line_match.start()
                    line_actual_end = start + caption_line_match.end()
                    
                    # Find the actual line break
                    next_linebreak = content.find('\n', line_actual_end)
                    if next_linebreak != -1:
                        line_actual_end = next_linebreak + 1
                    
                    # Remove the caption line
                    content = content[:line_start] + content[line_actual_end:]
                    stats['removed_dups'] += 1
                    
                    # Insert figure HTML at the same position
                    content = content[:line_start] + figure_html + '\n' + content[line_start:]
                else:
                    # Insert after the paragraph
                    para_end = content.find('\n\n', start)
                    if para_end == -1:
                        para_end = line_end
                    content = content[:para_end] + '\n' + figure_html + content[para_end:]
    
    # Write back
    with open(md_file, 'w') as f:
        f.write(content)
    
    return stats

def main():
    print("="*70)
    print("COMPREHENSIVE FIGURE INTEGRATION v2")
    print("="*70 + "\n")
    
    # Create public dir
    os.makedirs(PUBLIC_DIR, exist_ok=True)
    
    # Step 1: Extract all figure captions
    print("1. Extracting figure captions from Word document...")
    all_figures = extract_all_figure_captions()
    total_figs = sum(len(figs) for figs in all_figures.values())
    print(f"   Found {total_figs} figures across {len(all_figures)} chapters\n")
    
    # Step 2: Scan available images
    print("2. Scanning available images...")
    available_images = scan_available_images()
    total_imgs = sum(len(imgs) for imgs in available_images.values())
    print(f"   Found {total_imgs} images\n")
    
    # Step 3: Process each chapter
    print("3. Processing chapters...")
    print("="*70)
    
    total_stats = {'added': 0, 'placeholder': 0, 'removed_dups': 0}
    
    for chapter_num in sorted(all_figures.keys()):
        stats = process_chapter(chapter_num, all_figures, available_images, PUBLIC_DIR)
        if stats:
            total_stats['added'] += stats['added']
            total_stats['placeholder'] += stats['placeholder']
            total_stats['removed_dups'] += stats['removed_dups']
            
            print(f"Chapter {chapter_num}:")
            print(f"  ✓ {stats['added']} figures with images")
            print(f"  ⚠ {stats['placeholder']} placeholders")
            print(f"  🗑 {stats['removed_dups']} duplicate captions removed")
    
    print("="*70)
    
    # Step 4: Regenerate Astro pages
    print("\n4. Regenerating Astro pages...")
    import subprocess
    result = subprocess.run(['python3', 'integrate_chapters.py'], 
                          capture_output=True, text=True)
    print(result.stdout)
    
    print("\n" + "="*70)
    print("✅ COMPLETE")
    print("="*70)
    print(f"\nTotal: {total_stats['added']}")
    print(f"  ✓ Figures with images: {total_stats['added']}")
    print(f"  ⚠ Placeholders: {total_stats['placeholder']}")
    print(f"  🗑 Duplicates removed: {total_stats['removed_dups']}")
    print("\nAll figures now open full-size when clicked!")
    print("="*70)

if __name__ == "__main__":
    main()
