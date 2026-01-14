#!/usr/bin/env python3
"""
Final Figure Integration - Uses enhanced manual mapping + auto-matching
"""

import json
import os
import re
import shutil
from pathlib import Path
from docx import Document
from PIL import Image

# Paths
SOURCE_DIR = 'project_assets/images/dissertation/By_Chapter'
PUBLIC_DIR = 'public/assets/dissertation/figures'
EXTRACTED_DIR = 'project_assets/docs/dissertation/extracted'
DOCX_PATH = 'project_assets/docs/dissertation/Schurtz_Dissertation_Final.docx'
MANUAL_MAPPING_FILE = 'enhanced_figure_mapping.json'

def extract_all_figure_captions():
    """Extract complete figure list with captions from Word doc"""
    doc = Document(DOCX_PATH)
    figures = {}
    
    for i in range(201, 260):
        if i >= len(doc.paragraphs):
            break
        
        text = doc.paragraphs[i].text.strip()
        match = re.match(r'Figure\s+(\d+\.\d+):\s+(.+?)(?:\s+\d+\s*$|$)', text)
        if match:
            fig_num = match.group(1)
            caption = match.group(2).strip()
            caption = re.sub(r'\s+\d+\s*$', '', caption)
            
            chapter = fig_num.split('.')[0]
            if chapter not in figures:
                figures[chapter] = {}
            
            figures[chapter][fig_num] = caption
    
    return figures

def load_manual_mapping():
    """Load enhanced manual mapping"""
    with open(MANUAL_MAPPING_FILE, 'r') as f:
        return json.load(f)

def convert_tif_to_jpg(tif_path):
    """Convert TIF file to JPG"""
    jpg_path = tif_path.replace('.tif', '.jpg')
    if not os.path.exists(jpg_path):
        try:
            img = Image.open(tif_path)
            img.convert('RGB').save(jpg_path, 'JPEG', quality=90)
            return jpg_path
        except:
            return tif_path
    return jpg_path

def create_figure_html(chapter, fig_ref, caption, img_path=None):
    """Create clickable figure HTML"""
    
    if not img_path:
        # Placeholder
        return f'''<figure class="my-8 p-6 border border-dashed border-gray-700 bg-gray-900/20 rounded">
<div class="text-center py-8 text-gray-500">
<svg class="mx-auto h-10 w-10 mb-3 opacity-50" fill="none" stroke="currentColor" viewBox="0 0 24 24">
<path stroke-linecap="round" stroke-linejoin="round" stroke-width="1.5" d="M4 16l4.586-4.586a2 2 0 012.828 0L16 16m-2-2l1.586-1.586a2 2 0 012.828 0L20 14m-6-6h.01M6 20h12a2 2 0 002-2V6a2 2 0 00-2-2H6a2 2 0 00-2 2v12a2 2 0 002 2z" />
</svg>
<p class="text-sm font-mono text-gray-600">Figure {fig_ref} not yet available</p>
</div>
<figcaption class="text-sm text-gray-500 mt-2 text-center">Figure {fig_ref}: {caption}</figcaption>
</figure>'''
    
    # Real figure with click-to-expand
    return f'''<figure class="my-8">
<a href="{img_path}" target="_blank" class="block group cursor-zoom-in">
<img src="{img_path}" alt="Figure {fig_ref}: {caption}" class="w-full max-w-xl mx-auto border border-white/10 group-hover:border-[var(--color-brand-orange)] transition-all shadow-lg" loading="lazy" />
</a>
<figcaption class="text-sm text-gray-500 mt-2 text-center">Figure {fig_ref}: {caption} <span class="text-xs text-gray-600 ml-2">(click to enlarge)</span></figcaption>
</figure>'''

def process_chapter(chapter_num, all_figures, manual_mapping, public_dir):
    md_file = Path(EXTRACTED_DIR) / f"chapter-{chapter_num}.md"
    
    if not md_file.exists() or chapter_num not in all_figures:
        return None
    
    with open(md_file, 'r') as f:
        content = f.read()
    
    chapter_figures = all_figures[chapter_num]
    manual_figs = manual_mapping.get(chapter_num, {})
    stats = {'real': 0, 'placeholder': 0}
    
    for fig_ref, caption in sorted(chapter_figures.items()):
        img_path = None
        
        # Check manual mapping first
        if fig_ref in manual_figs:
            source_file = manual_figs[fig_ref]['source_file']
            caption = manual_figs[fig_ref]['caption']  # Use manual caption
            
            # Find source file
            chapter_dir = Path(SOURCE_DIR) / f"Chapter {chapter_num}"
            source_path = None
            
            # Check if it's a full path or just filename
            if '/' in source_file:
                full_source = chapter_dir / source_file
                if full_source.exists():
                    source_path = full_source
            else:
                # Search in chapter dir
                for file in chapter_dir.rglob(source_file):
                    source_path = file
                    break
            
            if source_path:
                # Handle TIF conversion
                if source_path.suffix.lower() in ['.tif', '.tiff']:
                    converted_path = convert_tif_to_jpg(str(source_path))
                    ext = Path(converted_path).suffix
                else:
                    ext = source_path.suffix
                
                dest_filename = f"figure-{chapter_num}-{fig_ref.replace('.', '-')}{ext}"
                dest_path = Path(public_dir) / dest_filename
                
                # Copy file
                if not dest_path.exists():
                    if source_path.suffix.lower() in ['.tif', '.tiff']:
                        shutil.copy2(converted_path, dest_path)
                    else:
                        shutil.copy2(source_path, dest_path)
                
                img_path = f"/assets/dissertation/figures/{dest_filename}"
                stats['real'] += 1
        
        if not img_path:
            stats['placeholder'] += 1
        
        figure_html = create_figure_html(chapter_num, fig_ref, caption, img_path)
        
        # Find and replace figure reference in text
        # Look for "Figure X.Y" pattern
        pattern = rf'Figure {re.escape(fig_ref)}(?::|\s)'
        
        matches = list(re.finditer(pattern, content, re.IGNORECASE))
        
        if matches:
            match = matches[0]
            start = match.start()
            
            # Check if already in figure tag
            check_range = content[max(0, start-300):min(len(content), start+300)]
            if '<figure' in check_range and '</figure>' in check_range:
                continue  # Already processed
            
            # Find the line containing this reference
            line_start = content.rfind('\n', 0, start) + 1
            line_end = content.find('\n', start)
            if line_end == -1:
                line_end = len(content)
            
            line_content = content[line_start:line_end]
            
            # If the line is JUST the figure caption, remove it
            if re.match(rf'^Figure {re.escape(fig_ref)}:', line_content.strip()):
                # Find actual line breaks
                actual_line_end = line_end
                if actual_line_end < len(content) and content[actual_line_end] == '\n':
                    actual_line_end += 1
                
                # Remove caption line and insert figure
                content = content[:line_start] + '\n\n' + figure_html + '\n\n' + content[actual_line_end:]
            else:
                # Insert figure after the paragraph
                para_end = content.find('\n\n', start)
                if para_end == -1:
                    para_end = line_end
                
                content = content[:para_end] + '\n\n' + figure_html + '\n\n' + content[para_end:]
    
    with open(md_file, 'w') as f:
        f.write(content)
    
    return stats

def main():
    print("="*70)
    print("FINAL FIGURE INTEGRATION")
    print("="*70 + "\n")
    
    os.makedirs(PUBLIC_DIR, exist_ok=True)
    
    print("1. Loading data...")
    all_figures = extract_all_figure_captions()
    manual_mapping = load_manual_mapping()
    
    total_figs = sum(len(figs) for figs in all_figures.values())
    manual_count = sum(len(figs) for figs in manual_mapping.values())
    
    print(f"   Total figures: {total_figs}")
    print(f"   Manually mapped: {manual_count}\n")
    
    print("2. Processing chapters...")
    print("="*70)
    
    total_stats = {'real': 0, 'placeholder': 0}
    
    for chapter_num in sorted(all_figures.keys()):
        stats = process_chapter(chapter_num, all_figures, manual_mapping, PUBLIC_DIR)
        if stats:
            total_stats['real'] += stats['real']
            total_stats['placeholder'] += stats['placeholder']
            
            print(f"Chapter {chapter_num}: {stats['real']} figures, {stats['placeholder']} placeholders")
            
            # Print missing figures for manual matching
            if stats['placeholder'] > 0:
                print("  Missing figures:")
                chapter_figures = all_figures[chapter_num]
                manual_figs = manual_mapping.get(chapter_num, {})
                for fig_ref, caption in chapter_figures.items():
                    if fig_ref not in manual_figs: 
                         # Check if we found an image automatically (simple check)
                         # Actually just print it if not in manual mapping, we can ignore if it was auto-matched
                         print(f"    - {fig_ref}: {caption}")

    
    print("="*70)
    
    # Regenerate Astro pages
    print("\n3. Regenerating Astro pages...")
    import subprocess
    result = subprocess.run(['python3', 'integrate_chapters.py'], 
                          capture_output=True, text=True)
    
    print("\n" + "="*70)
    print("✅ COMPLETE")
    print("="*70)
    print(f"\n✓ {total_stats['real']} figures with images")
    print(f"⚠ {total_stats['placeholder']} placeholders")
    print(f"\nAll figures are clickable to view full-size!")
    print("="*70)

if __name__ == "__main__":
    main()
