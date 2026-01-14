#!/usr/bin/env python3
"""
Dissertation Content Extractor v3 for marshallschurtz.com

Based on document analysis:
- Chapters are Heading 1 style paragraphs starting with ":"
- Chapter boundaries defined by paragraph indices
"""

import json
import os
import re
from docx import Document

# Configuration
DOCX_PATH = 'project_assets/docs/dissertation/Schurtz_Dissertation_Final.docx'
OUTPUT_DIR = 'project_assets/docs/dissertation/extracted'

# Exact chapter positions from document analysis
CHAPTERS = [
    {"number": "1", "title": "Introduction & Geographical Background", "start_idx": 261, "end_idx": 298},
    {"number": "2", "title": "Historical Background of Sidekan", "start_idx": 298, "end_idx": 461},
    {"number": "3", "title": "Archaeological Background of Sidekan and Soran", "start_idx": 461, "end_idx": 518},
    {"number": "4", "title": "Excavations of Gund-i Topzawa, Ghaberstan-i Topzawa, Sidekan Bank", "start_idx": 518, "end_idx": 738},
    {"number": "5", "title": "Survey of the Sidekan Subdistrict", "start_idx": 738, "end_idx": 876},
    {"number": "6", "title": "The Landscape and Settlement Patterns of the Sidekan Subdistrict", "start_idx": 876, "end_idx": 977},
    {"number": "7", "title": "Conclusion - The Character and Origin of Muṣaṣir", "start_idx": 977, "end_idx": 1088},
]


def get_heading_level(para):
    """Get heading level from paragraph style"""
    style_name = para.style.name if para.style else ""
    if style_name == "Heading 1":
        return 1
    elif style_name == "Heading 2":
        return 2
    elif style_name == "Heading 3":
        return 3
    elif style_name == "Heading 4":
        return 4
    return 0


def clean_text(text):
    """Clean up text for markdown output"""
    text = ' '.join(text.split())
    return text.strip()


def detect_figure_references(text):
    """Find figure references in text"""
    # Match Figure X.X patterns
    matches = re.findall(r'[Ff]igure\s+(\d+[.\-]\d+)', text)
    return matches


def extract_chapter(doc, chapter_info):
    """Extract content for a single chapter"""
    
    chapter_data = {
        "number": chapter_info["number"],
        "title": chapter_info["title"],
        "content": [],
        "figures": [],
        "word_count": 0
    }
    
    start_idx = chapter_info["start_idx"] + 1  # Skip the chapter heading itself
    end_idx = chapter_info["end_idx"]
    
    for i in range(start_idx, end_idx):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        if not text:
            continue
        
        style = para.style.name if para.style else "Normal"
        heading_level = get_heading_level(para)
        
        cleaned = clean_text(text)
        
        if cleaned and len(cleaned) > 1:
            content_type = "heading" if heading_level > 0 else "paragraph"
            
            chapter_data["content"].append({
                "type": content_type,
                "level": heading_level,
                "text": cleaned,
                "style": style
            })
            chapter_data["word_count"] += len(cleaned.split())
            
            # Track figure references
            figs = detect_figure_references(cleaned)
            chapter_data["figures"].extend(figs)
    
    return chapter_data


def convert_to_markdown(chapter_data):
    """Convert chapter content to Markdown format"""
    md_lines = []
    
    # Frontmatter
    md_lines.append("---")
    md_lines.append(f"chapter: {chapter_data['number']}")
    md_lines.append(f"title: \"{chapter_data['title']}\"")
    md_lines.append(f"word_count: {chapter_data['word_count']}")
    unique_figs = list(set(chapter_data['figures']))
    md_lines.append(f"figures: {len(unique_figs)}")
    md_lines.append("---")
    md_lines.append("")
    
    # Chapter heading
    md_lines.append(f"# Chapter {chapter_data['number']}: {chapter_data['title']}")
    md_lines.append("")
    
    for item in chapter_data["content"]:
        text = item["text"]
        item_type = item["type"]
        level = item.get("level", 0)
        
        if item_type == "heading" and level > 0:
            # h2 for Heading 2, h3 for Heading 3, etc.
            prefix = "#" * (level + 1)
            md_lines.append("")
            md_lines.append(f"{prefix} {text}")
            md_lines.append("")
        else:
            md_lines.append(text)
            md_lines.append("")
    
    return "\n".join(md_lines)


def main():
    # Create output directory
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    print(f"Loading document: {DOCX_PATH}")
    doc = Document(DOCX_PATH)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    extracted_chapters = []
    
    print("\n" + "="*60)
    print("Extracting chapters...")
    print("="*60)
    
    for chapter_info in CHAPTERS:
        chapter_data = extract_chapter(doc, chapter_info)
        extracted_chapters.append(chapter_data)
        
        unique_figs = len(set(chapter_data["figures"]))
        print(f"Chapter {chapter_data['number']}: {len(chapter_data['content'])} paragraphs, {chapter_data['word_count']:,} words, {unique_figs} figures")
    
    # Write markdown files
    print("\n" + "="*60)
    print("Writing Markdown files...")
    print("="*60)
    
    for chapter_data in extracted_chapters:
        md_content = convert_to_markdown(chapter_data)
        output_path = os.path.join(OUTPUT_DIR, f"chapter-{chapter_data['number']}.md")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        print(f"  ✓ chapter-{chapter_data['number']}.md ({chapter_data['word_count']:,} words)")
    
    # Write summary JSON
    summary = {
        "chapters": [
            {
                "number": ch["number"],
                "title": ch["title"],
                "word_count": ch["word_count"],
                "paragraph_count": len(ch["content"]),
                "figures_referenced": list(set(ch["figures"])),
                "md_file": f"chapter-{ch['number']}.md"
            }
            for ch in extracted_chapters
        ],
        "total_words": sum(ch["word_count"] for ch in extracted_chapters)
    }
    
    summary_path = os.path.join(OUTPUT_DIR, "extraction_summary.json")
    with open(summary_path, 'w', encoding='utf-8') as f:
        json.dump(summary, f, indent=2)
    print(f"  ✓ extraction_summary.json")
    
    # Summary stats
    print("\n" + "="*60)
    print("EXTRACTION COMPLETE")
    print("="*60)
    print(f"\nOutput: {OUTPUT_DIR}/")
    print(f"Chapters: {len(extracted_chapters)}")
    print(f"Total words: {summary['total_words']:,}")
    
    total_figures = sum(len(set(ch["figures"])) for ch in extracted_chapters)
    print(f"Total figure references: {total_figures}")
    
    print("\n" + "-"*60)
    print("FILES CREATED:")
    for ch in extracted_chapters:
        print(f"  - chapter-{ch['number']}.md: Chapter {ch['number']} ({ch['word_count']:,} words)")
    print("  - extraction_summary.json")
    print("-"*60)


if __name__ == "__main__":
    main()
